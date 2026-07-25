"""
Tests for docx_export.py.

Includes a round-trip test that feeds REAL writer_node output into
render_report_to_docx — this is what actually caught the underscore-italic
bug during development (the empty-section marker leaking literal
underscores into the rendered document), so it stays as a permanent
regression guard rather than only testing docx_export in isolation with
hand-built fixtures.

Note on history: an earlier version of this file asserted a native Word
TOC *field* (`"TOC" in body_xml`, `"updateFields" in settings_xml`). That
approach was replaced after empirically rendering the doc and finding the
field shows literally "Right-click and choose Update Field..." on headless
conversion rather than real entries — see docx_export.py's module
docstring. The contents-list tests below assert the fixed (manual
bookmarked hyperlink) behavior instead.
"""
from __future__ import annotations

import io
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from app.graph.nodes import writer as writer_module
from app.models.schemas import (
    ReportSection,
    ResearchAngle,
    ResearchMethod,
    SourceDoc,
    SubQuestion,
    VerificationStatus,
    VerifiedClaim,
)
from app.models.schemas import Report as ReportModel
from app.services.docx_export import EMPTY_SECTION_MARKER, render_report_to_docx

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _sample_report(**overrides) -> ReportModel:
    doc_a = SourceDoc(sub_question_id="sq1", url="https://reuters.com/a", title="Reuters Article", snippet="s")
    doc_b = SourceDoc(sub_question_id="sq1", url="https://apnews.com/b", title="AP Article", snippet="s")
    defaults = dict(
        session_id="s1",
        title="Test Report Title",
        executive_summary="This is the executive summary.",
        sections=[
            ReportSection(
                heading="First sub-question",
                content="- A corroborated claim [1][2]",
                cited_claim_ids=[],
            ),
            ReportSection(
                heading="Second sub-question",
                content="- **Disputed:** A disputed claim [1], contradicted by [2]. Notes here.",
                cited_claim_ids=[],
            ),
        ],
        citations={doc_a.id: doc_a, doc_b.id: doc_b},
        limitations=["Some open question"],
    )
    defaults.update(overrides)
    return ReportModel(**defaults)


def _open_docx(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def _all_text(document: Document) -> str:
    return "\n".join(p.text for p in document.paragraphs)


def _bookmark_ids(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    return [el.get(f"{W_NS}id") for el in root.iter(f"{W_NS}bookmarkStart")]


# ---- Basic structure ----

def test_render_produces_valid_docx_bytes():
    report = _sample_report()
    data = render_report_to_docx(report)
    assert data[:2] == b"PK"  # docx is a zip archive
    _open_docx(data)  # should open without raising


def test_title_is_rendered_as_heading():
    report = _sample_report()
    document = _open_docx(render_report_to_docx(report))
    heading_texts = [
        p.text for p in document.paragraphs
        if p.style.name.startswith("Title") or p.style.name.startswith("Heading")
    ]
    assert "Test Report Title" in heading_texts


def test_all_section_headings_present():
    report = _sample_report()
    document = _open_docx(render_report_to_docx(report))
    all_text = [p.text for p in document.paragraphs]
    assert "First sub-question" in all_text
    assert "Second sub-question" in all_text


def test_executive_summary_present():
    report = _sample_report()
    document = _open_docx(render_report_to_docx(report))
    assert "This is the executive summary." in _all_text(document)


def test_limitations_section_present():
    report = _sample_report()
    document = _open_docx(render_report_to_docx(report))
    all_text = [p.text for p in document.paragraphs]
    assert "Limitations" in all_text
    assert "Some open question" in all_text


def test_handles_empty_report_gracefully():
    report = _sample_report(sections=[], citations={}, limitations=[])
    document = _open_docx(render_report_to_docx(report))
    assert "References" in _all_text(document)


# ---- Disputed/unverified rendering ----

def test_disputed_claim_bold_and_text_both_present():
    report = _sample_report()
    document = _open_docx(render_report_to_docx(report))
    all_text = _all_text(document)
    assert "Disputed:" in all_text
    assert "A disputed claim" in all_text
    assert "Notes here." in all_text

    bold_found = any(
        "Disputed:" in run.text and run.bold
        for p in document.paragraphs
        for run in p.runs
    )
    assert bold_found


def test_unverified_marker_renders_without_asterisks():
    report = _sample_report(
        sections=[ReportSection(
            heading="H1", content="- Some claim [1] *(unverified — single source)*", cited_claim_ids=[],
        )],
    )
    document = _open_docx(render_report_to_docx(report))
    all_text = _all_text(document)
    assert "unverified" in all_text.lower()
    assert "*" not in all_text


# ---- The empty-section bug this was built to catch ----

def test_empty_section_marker_does_not_leak_underscores():
    report = _sample_report(sections=[ReportSection(heading="H1", content=EMPTY_SECTION_MARKER, cited_claim_ids=[])])
    document = _open_docx(render_report_to_docx(report))
    all_text = _all_text(document)
    assert "_No verified" not in all_text
    assert "No verified claims were found for this sub-question." in all_text


# ---- References / citations ----

def test_references_section_lists_all_citations_with_correct_numbering():
    report = _sample_report()
    document = _open_docx(render_report_to_docx(report))
    all_text = _all_text(document)
    assert "1. " in all_text and "Reuters Article" in all_text
    assert "2. " in all_text and "AP Article" in all_text


def test_reference_hyperlinks_point_to_correct_urls():
    report = _sample_report()
    data = render_report_to_docx(report)
    document = _open_docx(data)

    # python-docx doesn't expose w:hyperlink relationships via the high-level
    # API, so inspect the underlying XML relationships directly.
    rels = document.part.rels
    target_urls = {rel.target_ref for rel in rels.values() if rel.reltype.endswith("hyperlink")}
    assert "https://reuters.com/a" in target_urls
    assert "https://apnews.com/b" in target_urls


def test_citation_markers_are_internal_hyperlinks_to_bookmarks():
    report = _sample_report()
    data = render_report_to_docx(report)
    document = _open_docx(data)

    xml = document.element.xml
    assert 'w:anchor="ref1"' in xml
    assert 'w:anchor="ref2"' in xml
    assert 'w:name="ref1"' in xml
    assert 'w:name="ref2"' in xml


def test_quality_flags_shown_in_references():
    doc_a = SourceDoc(
        sub_question_id="sq1", url="https://quora.com/x", title="Q post", snippet="s",
        quality_flags=["content_farm_domain", "no_date"],
    )
    report = _sample_report(citations={doc_a.id: doc_a})
    document = _open_docx(render_report_to_docx(report))
    assert "content_farm_domain" in _all_text(document)


# ---- Contents list (replaces the old, broken native-TOC-field design) ----

def test_contents_list_has_internal_hyperlinks_to_every_heading():
    report = _sample_report()
    data = render_report_to_docx(report)
    document = _open_docx(data)
    xml = document.element.xml

    # Every heading must have a matching bookmark, and the contents list
    # (built before any heading, so it can link forward) must anchor to it.
    for anchor in ("sec-exec-summary", "sec-0", "sec-1", "sec-limitations", "sec-references"):
        assert f'w:name="{anchor}"' in xml
        assert f'w:anchor="{anchor}"' in xml


def test_contents_list_entries_appear_before_the_sections_they_link_to():
    report = _sample_report()
    document = _open_docx(render_report_to_docx(report))
    all_text = [p.text for p in document.paragraphs]

    toc_index = all_text.index("Table of Contents")
    first_section_heading_index = all_text.index("First sub-question")
    # The contents list entry for "First sub-question" appears once as a
    # link (before the real heading) and again as the heading itself.
    assert all_text.count("First sub-question") == 2
    assert toc_index < first_section_heading_index


def test_bookmark_ids_are_all_unique():
    doc_a, doc_b, doc_c = [
        SourceDoc(sub_question_id="sq", url=u, title="T", snippet="s")
        for u in ("https://a.com", "https://b.com", "https://c.com")
    ]
    sections = [
        ReportSection(heading=f"Heading {i}", content=f"- claim {i} [1]", cited_claim_ids=[])
        for i in range(5)
    ]
    report = _sample_report(
        sections=sections,
        citations={doc_a.id: doc_a, doc_b.id: doc_b, doc_c.id: doc_c},
        limitations=["some gap"],
    )
    data = render_report_to_docx(report)
    ids = _bookmark_ids(data)
    assert len(ids) == len(set(ids))
    assert len(ids) > 0


# ---- Round-trip against REAL writer_node output (caught the original bug) ----

@pytest.mark.asyncio
async def test_round_trip_from_real_writer_node_output(monkeypatch):
    sq_disputed = SubQuestion(
        parent_query="test", question_text="is X disputed?",
        angle=ResearchAngle.FACTUAL, research_method=ResearchMethod.WEB_SEARCH,
    )
    sq_empty = SubQuestion(
        parent_query="test", question_text="uncovered question",
        angle=ResearchAngle.FACTUAL, research_method=ResearchMethod.WEB_SEARCH,
    )
    doc_a = SourceDoc(sub_question_id=sq_disputed.id, url="https://a.com", title="Source A", snippet="s")
    doc_b = SourceDoc(sub_question_id=sq_disputed.id, url="https://b.com", title="Source B", snippet="s")

    disputed_claim = VerifiedClaim(
        source_doc_id=doc_a.id, sub_question_id=sq_disputed.id, claim_text="Disputed fact",
        confidence=0.3, supporting_excerpt="", verification_status=VerificationStatus.DISPUTED,
        corroborating_source_ids=[], contradicting_source_ids=[doc_b.id],
        critic_notes="sources disagree", adjusted_confidence=0.3,
    )

    class _NoLLM:
        is_configured = False

    monkeypatch.setattr(writer_module, "build_llm_client", lambda: _NoLLM())

    state = {
        "session_id": "s1", "original_query": "is X disputed?",
        "sub_questions": [sq_disputed, sq_empty], "verified_claims": [disputed_claim],
        "source_docs": [doc_a, doc_b], "open_questions": [sq_empty.question_text],
    }
    writer_result = await writer_module.writer_node(state)
    report = writer_result["report"]

    data = render_report_to_docx(report)
    document = _open_docx(data)
    all_text = _all_text(document)

    assert "Disputed fact" in all_text
    assert "Disputed:" in all_text
    assert "_No verified" not in all_text
    assert "No verified claims were found for this sub-question." in all_text
    assert "**" not in all_text
    ids = _bookmark_ids(data)
    assert len(ids) == len(set(ids))
