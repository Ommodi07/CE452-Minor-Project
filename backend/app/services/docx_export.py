"""
Converts a Report into a polished .docx using python-docx (not the docx-js
skill workflow — this runs as part of a live FastAPI endpoint, not a one-off
sandbox document, so a Python library with no Node runtime dependency is the
right tool for code embedded in the user's own backend). The same domain
knowledge from the docx skill still applies here, translated to
python-docx's API:

  - Headings use python-docx's BUILT-IN heading styles (`add_heading`, which
    applies the "Heading N" style) — the standard prerequisite for a
    document to have real navigable structure (Word's own Navigation Pane,
    a TOC field if the user inserts one later, etc), same reason docx-js
    requires `HeadingLevel.*`.
  - The table of contents is a MANUAL list of real internal hyperlinks
    (bookmarked headings + `w:anchor` links), not a native Word TOC field.
    This was verified empirically: a native TOC field (`add_toc` +
    `w:updateFields`) only populates its entries when Word/LibreOffice
    recalculates fields, which does NOT happen on headless conversion —
    rendering to PDF for verification showed literally "Right-click and
    choose Update Field..." instead of entries. Real end users opening this
    in Word would likely see it auto-update, but not everyone opens a
    generated report in desktop Word specifically (Google Docs import,
    LibreOffice, Word Online without auto-update) — a manual link list
    resolves correctly in every viewer immediately, no update step required.
  - Citations are real internal hyperlinks (`[N]` in a section body jumps to
    a bookmark on that reference's entry), and each reference entry is a
    real external hyperlink to its source url — not plain text that merely
    looks like a link.

Deliberately parses `ReportSection.content` (the markdown Writer already
produced) rather than re-deriving from raw VerifiedClaims, which the export
endpoint doesn't have access to (only the persisted Report). This is safe
because Writer only ever emits a small, fixed markdown vocabulary — bullet
lines, **bold**, *italic*, [N] citation markers, and one fixed
underscore-italic marker for empty sections — not general markdown, so a
small hand-rolled tokenizer is sufficient and far more robust here than
pulling in a general markdown parser. The empty-section marker is
special-cased (see EMPTY_SECTION_MARKER) rather than taught to the regex
tokenizer, since it's the one place Writer uses underscore- rather than
asterisk-delimited italics — this was caught by rendering real output and
seeing literal underscores leak through before the special-case was added.
"""
from __future__ import annotations

import io
import itertools
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.models.schemas import Report

_HYPERLINK_COLOR = "0563C1"  # standard Word hyperlink blue

_INLINE_TOKEN_RE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|\*(?P<italic>[^*\n]+?)\*"
    r"|\[(?P<citation>\d+)\]"
)

# The one place Writer uses underscore- rather than asterisk-delimited
# italics (see writer.py's _render_section) — special-cased rather than
# taught to _INLINE_TOKEN_RE so that regex stays narrowly asterisk-based
# and doesn't risk misfiring on legitimate underscores elsewhere (URLs,
# snake_case terms in a claim's text, etc).
EMPTY_SECTION_MARKER = "_No verified claims were found for this sub-question._"


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    # bookmarkStart must come after pPr (if present), not before.
    insert_at = 1 if len(p) and p[0].tag == qn("w:pPr") else 0
    p.insert(insert_at, start)
    p.append(end)


def _add_hyperlink_run(paragraph, text: str, *, url: str | None = None, anchor: str | None = None):
    """Add a real Word hyperlink run — external (`url`) or internal (`anchor`, jumps to a bookmark)."""
    hyperlink = OxmlElement("w:hyperlink")
    if url is not None:
        r_id = paragraph.part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink.set(qn("r:id"), r_id)
    else:
        hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _HYPERLINK_COLOR)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_contents_entry(document, text: str, anchor: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    _add_hyperlink_run(p, text, anchor=anchor)


def _write_inline_markdown(paragraph, text: str) -> None:
    """Render our controlled markdown subset (**bold**, *italic*, [N]) as styled/linked runs."""
    pos = 0
    for match in _INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        if match.group("bold") is not None:
            paragraph.add_run(match.group("bold")).bold = True
        elif match.group("italic") is not None:
            paragraph.add_run(match.group("italic")).italic = True
        elif match.group("citation") is not None:
            n = match.group("citation")
            _add_hyperlink_run(paragraph, f"[{n}]", anchor=f"ref{n}")

        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_report_to_docx(report: Report) -> bytes:
    document = Document()

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Unique per document (OOXML requires this), decoupled from any
    # user-visible numbering (citation numbers, section order, etc).
    bookmark_ids = itertools.count(1)

    document.add_heading(report.title, level=0)
    meta = document.add_paragraph()
    meta_run = meta.add_run(f"Generated {report.generated_at.strftime('%B %d, %Y')}")
    meta_run.italic = True
    meta_run.font.size = Pt(10)

    # Bookmark anchors assigned up front so the contents list (built first)
    # can link forward to headings that appear later in the document.
    exec_summary_anchor = "sec-exec-summary"
    section_anchors = [f"sec-{i}" for i in range(len(report.sections))]
    limitations_anchor = "sec-limitations"
    references_anchor = "sec-references"

    document.add_heading("Table of Contents", level=1)
    _add_contents_entry(document, "Executive Summary", exec_summary_anchor)
    for report_section, anchor in zip(report.sections, section_anchors):
        _add_contents_entry(document, report_section.heading, anchor)
    if report.limitations:
        _add_contents_entry(document, "Limitations", limitations_anchor)
    _add_contents_entry(document, "References", references_anchor)
    document.add_page_break()

    heading = document.add_heading("Executive Summary", level=1)
    _add_bookmark(heading, exec_summary_anchor, next(bookmark_ids))
    document.add_paragraph(report.executive_summary)

    for report_section, anchor in zip(report.sections, section_anchors):
        heading = document.add_heading(report_section.heading, level=1)
        _add_bookmark(heading, anchor, next(bookmark_ids))

        content = report_section.content.strip()
        if content == EMPTY_SECTION_MARKER:
            p = document.add_paragraph()
            p.add_run(content.strip("_")).italic = True
            continue

        for line in content.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("- "):
                p = document.add_paragraph(style="List Bullet")
                _write_inline_markdown(p, line[2:])
            else:
                p = document.add_paragraph()
                _write_inline_markdown(p, line)

    if report.limitations:
        heading = document.add_heading("Limitations", level=1)
        _add_bookmark(heading, limitations_anchor, next(bookmark_ids))
        for item in report.limitations:
            document.add_paragraph(item, style="List Bullet")

    document.add_page_break()
    heading = document.add_heading("References", level=1)
    _add_bookmark(heading, references_anchor, next(bookmark_ids))
    for n, doc in enumerate(report.citations.values(), start=1):
        p = document.add_paragraph()
        _add_bookmark(p, f"ref{n}", next(bookmark_ids))
        p.add_run(f"{n}. ").bold = True
        _add_hyperlink_run(p, doc.title, url=doc.url)
        if doc.quality_flags:
            flag_run = p.add_run(f"  ({', '.join(doc.quality_flags)})")
            flag_run.italic = True
            flag_run.font.size = Pt(9)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
